"""
Coal Consumption & Environmental Performance Indicators (EPI) Extractor —
pulls plant-level monthly figures from EMD's monthly environmental report,
which comes in four source formats:

PDF, old "Major Environmental Performance Indicators" style (2 pages,
  page-1 title has NO "EMD Flash Report" text) —

  Page 1 — "Major Environmental Performance Indicators (EPIs)" table, one
    3-row-header block per parameter (Sp. CO2 Emission / Sp. Water Consumption
    / Sp. PM Emission), each with exactly 6 data rows (BSP, DSP, RSP, BSL, ISP,
    SAIL) in that fixed order. Column layout (which month columns exist, and
    whether "Target 2026-27" / "Actual <month>" split into Existing
    Calculation# / Additional Emission* sub-columns) varies month to month —
    columns are located by their own header text/position, not fixed indices.
    We only want the "Existing Calculation" figure where a split exists (the
    WSA CO2 baseline methodology), matching the other two params which never
    split. Located by text position (`page.get_text("words")`), not by
    PyMuPDF's `find_tables()` — table auto-detection was found to be unreliable
    across report months here (drops plant labels non-deterministically,
    splits one logical table into inconsistent fragments across sample months).

  Page 2 — "Consumption of Coking Coal and CDI Coal" table, left ('000 T
    quantity) half only. Each plant contributes one row per month it appears
    in the report (current month, plus running FY-cumulative rows for months
    after April) - we match on the exact current-month row label (e.g.
    "May'26", not "Apr-May'26") to avoid picking up the cumulative row. Same
    table, same extract_page2_coal(), is also used by the "EMD Flash Report"
    PDF's page 2 and by a standalone "Coal OMI" PDF that carries only this
    table as its own page 1 - extract_pdf() locates it by title text
    ("Consumption of Coking Coal"), not a fixed page index, for exactly
    that reason.

DOCX, "EMD Flash Report" .docx (seen starting Jul'26) — a slimmer
  preliminary release: a single table with Specific Water Consumption /
  Specific CO2 Emission (each Target / <month> / FY-cumulative sub-columns)
  and BF Slag Utilisation (not currently stored). No Coal Consumption table
  at all (that's a separate .pdf attachment for this format, if sent).
  Columns are located by header text via python-docx's table grid (merged
  header cells repeat their text across every column they span), not fixed
  indices, for the same month-to-month resilience as the old PDF path.

PDF, "EMD Flash Report" style (2 pages, page-1 title HAS "EMD Flash Report"
  text) — the same Flash Report as the .docx above, just delivered as a PDF
  (e.g. one plant row "BSP 2.70 2.74 2.74 2.45 2.51 2.53 92 94" — Water
  [Target, month, cum], CO2 [Target, month, cum], BF Slag [month, cum], in
  that fixed left-to-right order; PyMuPDF renders one token per line here,
  so extract_page1_flash() reads 8 consecutive numeric lines after each
  plant-label line rather than using word-position matching). Page 2 is the
  same Coal Consumption table as the old PDF style. The report's own
  "EMD Flash Report for <Month>, <Year>" page-1 title is checked against the
  selected report_month so a wrong month selection fails clearly instead of
  silently reading the right numbers under the wrong month.

  Because "Major Environmental Performance Indicators (EPIs)" appears as a
  section heading in BOTH PDF styles, extract_pdf() distinguishes them by
  the presence of "EMD Flash Report" on that same page, not by title alone.

XLSX, "Major EPIs <Mon>'<YY>.xlsx" workbook — see extract_xlsx()'s own
  docstring for the row/column layout. Reachable the same way as the other
  three formats via extract_report()/extract_pdf()'s callers (e.g.
  api_coal_co2_techno.py's /preview and /insert), not just the standalone
  load_xlsx() script entry point.

Every format but the old PDF carries no Sp. PM Emission (the old PDF always
has it) — enviro's "pm" key is simply absent then, and neither format but
the old/Flash PDF carries a Coal Consumption table, and
plant_techno_json()/plant_till_techno_json() treat every param key as
optional so whatever a given report doesn't carry is left unset (not
overwritten to None) rather than erroring. Likewise "coal" is {} for any
report with no Coal Consumption table (e.g. a standalone "Coal OMI" PDF,
which has that table as its own page 1 - see extract_pdf()'s
page-detection-by-title, not fixed index).

Values land in techno_data (unit='General', techno_json["month"]) via
db.merge_upsert_techno_data — plant-level only; SAIL is intentionally never
written here. The at-a-glance / major-techno pages compute SAIL for these
params as a Crude-Steel-weighted average across plants (see BF_SAIL_SPECS's
"cs"-weighted entries in page_techno.py), matching how "Specific Energy
Consumption" already works, rather than trusting the report's own SAIL row
(which the report computes with EMD's own, not necessarily identical,
weighting).

Run as a script to (re-)load a folder of these reports (PDF and/or .docx):
    python coal_co2_epi_extractor.py "D:\\opr-mis1\\Report_format\\Coal_co2"
"""
import re
import sys
from pathlib import Path

PLANTS = ["BSP", "DSP", "RSP", "BSL", "ISP"]

# techno_data["month"] key -> display unit, for the 7 new parameters this
# extractor is the sole source of.
ENVIRO_KEY_UNITS = {
    "sp_co2_emission":      "T/tcs",
    "sp_water_consumption": "m\u00b3/tcs",
    "sp_pm_emission":       "kg/tcs",
}
COAL_KEY_UNITS = {
    "indigenous_pcc":     "'000 T",
    "indigenous_mcc":     "'000 T",
    "imported_hard_coal": "'000 T",
    "imported_soft_coal": "'000 T",
}

# Display param name (as used in generate_major_techno_from_db /
# generate_summary_te_table / techno_plan_fy target JSON) for each of the 3
# enviro parameters, in the fixed top-to-bottom order they appear in the PDF.
ENVIRO_PARAM_ORDER = [
    ("co2",   "Sp. CO2 Emission",         "sp_co2_emission"),
    ("water", "Sp. Water Consumption",    "sp_water_consumption"),
    ("pm",    "Sp. PM Emission",          "sp_pm_emission"),
]

_FNAME_RE = re.compile(r"^([A-Za-z]{3})'?(\d{2})\.pdf$", re.IGNORECASE)
_MONTH_NUM = {"jan": 1, "feb": 2, "mar": 3, "apr": 4, "may": 5, "jun": 6,
              "jul": 7, "aug": 8, "sep": 9, "oct": 10, "nov": 11, "dec": 12}


def report_month_from_filename(fname: str):
    """"Apr'26.pdf" -> ("2026-04", "Apr'26")"""
    m = _FNAME_RE.match(fname)
    if not m:
        return None
    mon, yy = m.group(1).lower(), int(m.group(2))
    if mon not in _MONTH_NUM:
        return None
    year = 2000 + yy
    return f"{year}-{_MONTH_NUM[mon]:02d}", f"{mon.capitalize()}'{yy:02d}"


_MONTH_ABBR = {v: k for k, v in _MONTH_NUM.items()}


def mlabel_from_report_month(report_month: str) -> str:
    """"2026-04" -> "Apr'26" - the exact column-header text this report uses
    for its current-month column, independent of what the uploaded file
    happens to be named. Used so a user-selected report_month drives which
    column is read; if that column isn't on the page at all (wrong month
    selected for this file), extraction fails with a clear error rather
    than silently reading the wrong column."""
    year, mon_num = report_month.split("-")
    abbr = _MONTH_ABBR[int(mon_num)]
    return f"{abbr.capitalize()}'{year[-2:]}"


def till_mlabel_from_report_month(report_month: str) -> str:
    """"2026-07" -> "Apr'26-Jul'26" - the .docx Flash Report's FY-cumulative
    column header text (April of the FY through the report month). Only
    used for the .docx path (extract_docx) - the PDF report's page 1 has no
    such cumulative column. April itself has no distinct range observed in
    any sample report; callers treat a miss here as "no cumulative column
    for this report" rather than an error, so a wrong guess here just means
    till_month stays unpopulated instead of breaking extraction."""
    year, mon_num = report_month.split("-")
    mon_num = int(mon_num)
    fy_year = int(year) if mon_num >= 4 else int(year) - 1
    if mon_num == 4:
        return mlabel_from_report_month(report_month)
    return f"Apr'{str(fy_year)[-2:]}-{mlabel_from_report_month(report_month)}"


def _norm(s: str) -> str:
    return s.replace("\u2019", "'")


def _words(page):
    return [(w[0], w[1], w[2], w[3], _norm(w[4])) for w in page.get_text("words")]


def _find(words, text, x_min=None, x_max=None, y_min=None, y_max=None):
    out = []
    for w in words:
        if w[4] != text:
            continue
        if x_min is not None and w[0] < x_min:
            continue
        if x_max is not None and w[0] > x_max:
            continue
        if y_min is not None and w[1] < y_min:
            continue
        if y_max is not None and w[1] > y_max:
            continue
        out.append(w)
    return out


def _col_center(words, label_variants, y_max=145):
    """x-center of the value column for a top-level header (identified by
    label_variants, e.g. the month string or "2026-27"). When that header
    has split into Existing-Calculation / Additional-Emission sub-columns
    (only ever seen for Sp. CO2 Emission, and only in some report months),
    return the Existing-Calculation sub-column's center instead - picked as
    the "Existing" sub-header word nearest this header's own center, since
    two different top-level headers can each have their own "Existing"
    sub-label only ~100pt apart and a fixed window risks grabbing the wrong
    one."""
    hits = []
    for lbl in label_variants:
        hits += _find(words, lbl, y_max=y_max)
    if not hits:
        raise ValueError(f"header label not found: {label_variants}")
    hw = sorted(hits, key=lambda w: w[1])[0]
    lx0, lx1, ly1 = hw[0], hw[2], hw[3]
    label_center = (lx0 + lx1) / 2
    exist_hits = [w for w in words if w[4] == "Existing" and ly1 <= w[1] <= ly1 + 40]
    if exist_hits:
        ew = min(exist_hits, key=lambda w: abs((w[0] + w[2]) / 2 - label_center))
        ew_center = (ew[0] + ew[2]) / 2
        if abs(ew_center - label_center) < 90:
            return ew_center
    return label_center


def _nearest_in_row(words, row_y0, row_y1, col_x, tol=40, x_max=None):
    cands = []
    for w in words:
        if not (row_y0 <= w[1] <= row_y1):
            continue
        if x_max is not None and w[0] > x_max:
            continue
        cx = (w[0] + w[2]) / 2
        if abs(cx - col_x) <= tol:
            cands.append((abs(cx - col_x), w[4]))
    if not cands:
        return None
    cands.sort(key=lambda t: t[0])
    return cands[0][1]


def extract_page1_enviro(page, mlabel) -> dict:
    """-> {"co2": {"month": {plant: val}, "target": {plant: val}}, "water": ..., "pm": ...}
    (plant keys include "SAIL" - callers should drop it; SAIL here is only
    useful as a cross-check against the plant sum, not for storage.)"""
    words = _words(page)

    label_words = [w for w in words if w[4] in (PLANTS + ["SAIL"]) and w[1] > 140 and w[0] < 350]
    label_words.sort(key=lambda w: w[1])
    if len(label_words) != 18:
        raise ValueError(f"expected 18 plant-row labels on the EPI page, got {len(label_words)}: "
                          f"{[(w[4], round(w[1], 1)) for w in label_words]}")

    target_col_x = _col_center(words, ["2026-27"])
    month_col_x = _col_center(words, [mlabel])

    blocks = {"co2": label_words[0:6], "water": label_words[6:12], "pm": label_words[12:18]}

    out = {}
    for key, row_words in blocks.items():
        month_vals, target_vals = {}, {}
        for w in row_words:
            plant = w[4]
            ry0, ry1 = w[1] - 3, w[3] + 3
            mv = _nearest_in_row(words, ry0, ry1, month_col_x)
            tv = _nearest_in_row(words, ry0, ry1, target_col_x)
            if mv is not None:
                try:
                    month_vals[plant] = float(mv)
                except ValueError:
                    pass
            if tv is not None:
                try:
                    target_vals[plant] = float(tv)
                except ValueError:
                    pass
        out[key] = {"month": month_vals, "target": target_vals}
    return out


def extract_page2_coal(page, mlabel) -> dict:
    """-> {plant: {"pcc":.., "mcc":.., "hard":.., "soft":..}} (Indigenous
    PCC/MCC, Imported Hard/Soft coking coal, '000 T) for the exact current
    month's row (not any FY-cumulative row also on the page). Includes SAIL
    (the report's own total) as a cross-check only."""
    words = _words(page)

    def hdr_x(text):
        hits = _find(words, text, x_max=400, y_max=110)
        if not hits:
            raise ValueError(f"coal-table header not found: {text}")
        w = hits[0]
        return (w[0] + w[2]) / 2

    pcc_x, mcc_x, hard_x, soft_x = hdr_x("PCC"), hdr_x("MCC"), hdr_x("Hard"), hdr_x("Soft")

    expected = PLANTS + ["SAIL"]
    month_row_hits = sorted(_find(words, mlabel, x_max=110), key=lambda w: w[1])
    if len(month_row_hits) != len(expected):
        raise ValueError(f"expected {len(expected)} '{mlabel}' rows on the coal page, got "
                          f"{len(month_row_hits)}: {[round(w[1], 1) for w in month_row_hits]}")

    out = {}
    for plant, mw in zip(expected, month_row_hits):
        ry0, ry1 = mw[1] - 3, mw[3] + 3
        vals = {}
        for name, cx in [("pcc", pcc_x), ("mcc", mcc_x), ("hard", hard_x), ("soft", soft_x)]:
            v = _nearest_in_row(words, ry0, ry1, cx, tol=25, x_max=400)
            if v is not None:
                try:
                    vals[name] = float(v)
                except ValueError:
                    pass
        out[plant] = vals
    return out


_FLASH_TITLE_RE = re.compile(r"EMD Flash Report\s+for\s*([A-Za-z]+)[,.]?\s*(\d{4})", re.I)


def _detect_flash_report_month(page_text: str):
    """"EMD Flash Report for July, 2026" -> "2026-07", or None if the title
    isn't found/parseable. Mirrors pdf_extractor_ssp's month-from-header
    approach (SSP's DPR) - each source report validates against its own
    stated month rather than trusting the uploaded filename or the
    user-selected month picker alone."""
    m = _FLASH_TITLE_RE.search(page_text[:400])
    if not m:
        return None
    mon_key = m.group(1).strip().lower()[:3]
    mon_num = _MONTH_NUM.get(mon_key)
    if not mon_num:
        return None
    return f"{int(m.group(2))}-{mon_num:02d}"


def extract_page1_flash(page) -> dict:
    """Extract Sp. CO2 Emission / Sp. Water Consumption from the "EMD Flash
    Report" PDF page's table - the PDF-page counterpart of extract_docx()
    for the same underlying report (see module docstring for the row
    layout). No Sp. PM Emission exists in this report, so the returned dict
    has no "pm" key. -> {"co2": {"month":.., "till_month":.., "target":..},
    "water": ...}"""
    lines = [ln.strip() for ln in (page.get_text() or "").splitlines() if ln.strip()]
    expected = PLANTS + ["SAIL"]

    idx = 0
    plant_vals = {}
    for plant in expected:
        while idx < len(lines) and lines[idx] != plant:
            idx += 1
        if idx >= len(lines):
            raise ValueError(f"'{plant}' row not found on the EMD Flash Report EPI table")
        idx += 1  # step past the plant-label line
        vals = []
        while len(vals) < 8 and idx < len(lines):
            try:
                vals.append(float(lines[idx].replace(",", "").rstrip("%")))
            except ValueError:
                break
            idx += 1
        if len(vals) != 8:
            raise ValueError(
                f"expected 8 numeric values (Water Target/Month/Cum, CO2 Target/Month/Cum, "
                f"Slag Month/Cum) for {plant} on the EMD Flash Report EPI table, got {len(vals)}")
        plant_vals[plant] = vals

    co2 = {"month": {}, "till_month": {}, "target": {}}
    water = {"month": {}, "till_month": {}, "target": {}}
    for plant, vals in plant_vals.items():
        water["target"][plant], water["month"][plant], water["till_month"][plant] = vals[0], vals[1], vals[2]
        co2["target"][plant], co2["month"][plant], co2["till_month"][plant] = vals[3], vals[4], vals[5]
        # vals[6], vals[7] = BF Slag Utilisation month/cum - not currently stored
    return {"co2": co2, "water": water}


def _find_page(doc, title_substr: str):
    """First page in doc whose text contains title_substr (case-insensitive),
    or None. Used instead of a fixed page index because some source reports
    (e.g. the standalone "Coal OMI" PDF, which carries only the Consumption
    of Coking Coal table as its own page 1) don't share the combined "Major
    EPI" report's page-1-enviro/page-2-coal layout."""
    for page in doc:
        if title_substr.lower() in (page.get_text() or "").lower():
            return page
    return None


def extract_pdf(pdf_path, report_month: str, mlabel: str) -> dict:
    """Full extraction for one month's report. Locates the EPI table and the
    Coal Consumption table each by their own title text rather than by fixed
    page index, so this also handles single-table PDFs that carry only one
    of the two (e.g. a standalone "Coal OMI" report has no EPI page) -
    whichever table is missing is simply left empty, same graceful handling
    as extract_docx's missing PM/Coal data. The EPI page is further
    dispatched to the old 18-row parser or the newer "EMD Flash Report"
    parser based on which one that specific page's own text matches (see
    module docstring) - a wrong month selection against a Flash Report PDF
    fails clearly via its own title date, same as extract_docx's mlabel check.
    -> {"enviro": {...} or {}, "coal": {...from extract_page2_coal} or {}}"""
    import fitz
    doc = fitz.open(pdf_path)
    enviro_page = _find_page(doc, "Major Environmental Performance Indicators")
    coal_page = _find_page(doc, "Consumption of Coking Coal")
    if enviro_page is None and coal_page is None:
        raise ValueError(
            "Could not find either the 'Major Environmental Performance Indicators' "
            "table or the 'Consumption of Coking Coal' table in this PDF."
        )

    enviro = {}
    if enviro_page is not None:
        page_text = enviro_page.get_text() or ""
        if "emd flash report" in page_text.lower():
            detected_month = _detect_flash_report_month(page_text)
            if detected_month and detected_month != report_month:
                raise ValueError(
                    f"This PDF's own header shows {detected_month}, but {report_month} "
                    f"was selected — please select the matching month."
                )
            enviro = extract_page1_flash(enviro_page)
        else:
            enviro = extract_page1_enviro(enviro_page, mlabel)

    return {
        "enviro": enviro,
        "coal": extract_page2_coal(coal_page, mlabel) if coal_page is not None else {},
    }


def extract_docx(docx_path, report_month: str, mlabel: str) -> dict:
    """Extract Sp. CO2 Emission / Sp. Water Consumption from the "EMD Flash
    Report" .docx format's single table, including its FY-cumulative
    ("Apr'YY-<mlabel>") column into "till_month". No Sp. PM Emission or Coal
    Consumption data exists in this report, so the returned dict has no
    "pm" key. -> {"co2": {"month": {...}, "till_month": {...}, "target": {...}}, "water": ...}"""
    import docx as _docx
    doc = _docx.Document(docx_path)
    if not doc.tables:
        raise ValueError("no tables found in the .docx report")
    table = doc.tables[0]
    n_rows, n_cols = len(table.rows), len(table.columns)
    if n_rows < 3:
        raise ValueError(f"expected a header + data table, got only {n_rows} row(s)")

    def cell(r, c):
        return _norm(table.cell(r, c).text.strip())

    hdr_param = [cell(0, c) for c in range(n_cols)]  # merged param-group header, repeated per spanned col
    hdr_sub   = [cell(1, c) for c in range(n_cols)]   # Target / <month> / <FY-cum> sub-header

    def cols_for(*needles):
        return [c for c, t in enumerate(hdr_param) if all(n in t for n in needles)]

    def sub_col(cols, label):
        for c in cols:
            if hdr_sub[c] == label:
                return c
        return None

    till_mlabel = till_mlabel_from_report_month(report_month)

    def block(param_needle, label):
        cols = cols_for(param_needle)
        if not cols:
            return {"month": {}, "till_month": {}, "target": {}}
        month_c = sub_col(cols, mlabel)
        if month_c is None:
            raise ValueError(
                f"'{mlabel}' column not found under {label} — "
                f"check the selected month matches the uploaded file")
        target_c = sub_col(cols, "Target")
        # Best-effort only — a miss just leaves till_month unpopulated for
        # this param rather than failing the whole extraction (see
        # till_mlabel_from_report_month's docstring on the April edge case).
        till_c = sub_col(cols, till_mlabel)

        def read_col(c, into):
            if c is None:
                return
            for r in range(2, n_rows):
                plant = cell(r, 0).upper()
                if plant not in PLANTS and plant != "SAIL":
                    continue
                v = cell(r, c)
                if v:
                    try:
                        into[plant] = float(v)
                    except ValueError:
                        pass

        month_vals, till_vals, target_vals = {}, {}, {}
        read_col(month_c, month_vals)
        read_col(till_c, till_vals)
        read_col(target_c, target_vals)
        return {"month": month_vals, "till_month": till_vals, "target": target_vals}

    return {
        "co2":   block("CO2", "Specific CO2 Emission"),
        "water": block("Water", "Specific Water Consumption"),
    }


def _xlsx_norm(v) -> str:
    return _norm(str(v)) if v is not None else ""


def _xlsx_find_col(ws, header_row: int, label: str):
    for col in range(1, ws.max_column + 1):
        if _xlsx_norm(ws.cell(row=header_row, column=col).value).strip() == label:
            return col
    return None


def _xlsx_block_starts(ws, plant_col: int = 2, label: str = "BSP") -> list:
    return [r for r in range(1, ws.max_row + 1)
            if _xlsx_norm(ws.cell(row=r, column=plant_col).value).strip() == label]


def extract_xlsx(xlsx_path, report_month: str, mlabel: str) -> dict:
    """Extract from the "Major EPIs <Mon>'<YY>.xlsx" workbook — a different
    source format from the PDF/.docx EMD reports above (single "Table 1"
    sheet, one column per period located by its own header text in row 4,
    e.g. "Jul'26"/"Jul'25"/"Apr.- Jul'26"/"Apr.- Jul'25" — column position
    isn't fixed, the sheet just grows another month-column each release,
    same resilience approach as the PDF/docx paths above).

    Unlike the PDF/docx reports (one file = one month), this workbook carries
    BOTH the current month AND the same month last year side by side in the
    same row (its own year-over-year comparison), plus each one's own
    Apr-to-month cumulative column — so a single call here extracts TWO
    report_months worth of data at once, keyed by report_month string.

    Row layout: 3 fixed parameter blocks (Sp. CO2 Emission, Sp. Water
    Consumption, Sp. PM Emission — same ENVIRO_PARAM_ORDER as the other
    formats), each block located by its own "BSP" cell in column B (the
    first plant row of that block) rather than a fixed row range, since one
    seen file has a quirk the others don't: the Water block's RSP and BSL
    rows are merged into a single "RSP BSL" cell with a 2-line ("<RSP>
    value>\\n<BSL value>") value in every data column instead of one row
    each — parse_block below handles both the normal 1-row-1-plant case and
    this 1-row-2-plants case generically (whichever labels are present),
    so it isn't tied to RSP+BSL specifically if a future month merges a
    different pair.

    No Target/Coal Consumption data is read here (this workbook's own
    "Target 2026-27" column exists but isn't wired to techno_plant_plan —
    only month/till_month is needed today).

    -> {report_month: {key: {"month":{plant:val}, "till_month":{plant:val}}},
        cply_report_month: {key: {...}}}
    (key = "co2"/"water"/"pm", same as every other extract_* here)."""
    import openpyxl
    wb = openpyxl.load_workbook(xlsx_path, data_only=True)
    ws = wb.worksheets[0]

    header_row = None
    for r in range(1, min(ws.max_row, 10) + 1):
        if any(_xlsx_norm(ws.cell(row=r, column=c).value).strip() == mlabel
               for c in range(1, ws.max_column + 1)):
            header_row = r
            break
    if header_row is None:
        raise ValueError(
            f"'{mlabel}' column not found in {xlsx_path} — "
            f"check the selected month matches the uploaded file")

    year, mon_num = report_month.split("-")
    cply_report_month = f"{int(year) - 1}-{mon_num}"
    cply_mlabel = mlabel_from_report_month(cply_report_month)
    # April has no "Apr.- Apr'YY" cumulative column in this workbook (same
    # edge case as till_mlabel_from_report_month above) — cumulative simply
    # isn't extracted for an April report_month.
    till_label      = f"Apr.- {mlabel}"      if int(mon_num) != 4 else None
    cply_till_label = f"Apr.- {cply_mlabel}" if int(mon_num) != 4 else None

    cols = {
        "month":      _xlsx_find_col(ws, header_row, mlabel),
        "cply_month": _xlsx_find_col(ws, header_row, cply_mlabel),
        "till_month": _xlsx_find_col(ws, header_row, till_label) if till_label else None,
        "cply_till":  _xlsx_find_col(ws, header_row, cply_till_label) if cply_till_label else None,
    }
    if cols["month"] is None:
        raise ValueError(
            f"'{mlabel}' column not found — check the selected month matches the uploaded file")
    if cols["cply_month"] is None:
        raise ValueError(f"comparable-prior-year column '{cply_mlabel}' not found in {xlsx_path}")

    block_starts = _xlsx_block_starts(ws)
    if len(block_starts) != len(ENVIRO_PARAM_ORDER):
        raise ValueError(
            f"expected {len(ENVIRO_PARAM_ORDER)} parameter blocks (one 'BSP' row each "
            f"in column B), found {len(block_starts)} in {xlsx_path}")

    plants_seq = PLANTS + ["SAIL"]

    def parse_block(start_row):
        out = {k: {} for k in cols}
        r, pi = start_row, 0
        while pi < len(plants_seq):
            label = _xlsx_norm(ws.cell(row=r, column=2).value).strip()
            two_ahead = plants_seq[pi:pi + 2]
            if len(two_ahead) == 2 and label == " ".join(two_ahead):
                plants_here = two_ahead
            elif label == plants_seq[pi]:
                plants_here = [plants_seq[pi]]
            else:
                raise ValueError(
                    f"unexpected plant label {label!r} at row {r} "
                    f"(expected {plants_seq[pi]!r}) in {xlsx_path}")
            for key, col in cols.items():
                if col is None:
                    continue
                raw = ws.cell(row=r, column=col).value
                if raw is None:
                    continue
                parts = [p.strip() for p in str(raw).split("\n")]
                if len(parts) != len(plants_here):
                    continue
                for plant, part in zip(plants_here, parts):
                    try:
                        out[key][plant] = float(part)
                    except ValueError:
                        pass
            pi += len(plants_here)
            r += 1
        return out

    enviro_cur, enviro_cply = {}, {}
    for (key, _label, _jk), start_row in zip(ENVIRO_PARAM_ORDER, block_starts):
        parsed = parse_block(start_row)
        enviro_cur[key]  = {"month": parsed["month"],      "till_month": parsed["till_month"]}
        enviro_cply[key] = {"month": parsed["cply_month"], "till_month": parsed["cply_till"]}

    return {report_month: enviro_cur, cply_report_month: enviro_cply}


def load_xlsx(xlsx_path, report_month: str, write: bool = True) -> dict:
    """(Re-)load one "Major EPIs <Mon>'<YY>.xlsx" workbook — see extract_xlsx
    for the format. Writes both the current report_month AND its
    comparable-prior-year month (both month + till_month figures) into
    techno_data, per plant, unit='General' — same target table/shape as
    load_folder's PDF/.docx writes, just two report_months from one file
    instead of one. No target/plan data is written (see extract_xlsx).
    Returns {report_month: {...}, cply_report_month: {...}} for inspection
    either way, mirroring load_folder's return shape."""
    sys.path.insert(0, str(Path(__file__).parent.parent))
    import db  # noqa: E402

    mlabel = mlabel_from_report_month(report_month)
    results = extract_xlsx(xlsx_path, report_month, mlabel)

    if write:
        for rm, enviro in results.items():
            for plant in PLANTS:
                month_json = plant_techno_json(enviro, {}, plant)
                till_json = plant_till_techno_json(enviro, plant)
                if month_json or till_json:
                    db.merge_upsert_techno_data(
                        plant, rm, "General", {"month": month_json, "till_month": till_json},
                        source_file=f"Coal_co2/{Path(xlsx_path).name}")

    return results


def extract_report(path, report_month: str, mlabel: str) -> dict:
    """Dispatches to the PDF, .docx or .xlsx extractor by file extension.
    -> {"enviro": {...}, "coal": {...}} — for .docx and .xlsx, "enviro" has
    no "pm"/"target" (.docx) or no "target" (.xlsx) keys and "coal" is
    always {}, since those reports don't carry those figures;
    plant_techno_json()/plant_till_techno_json() treat every key as
    optional. For .xlsx, extract_xlsx() also returns the comparable-prior-
    year month's figures (see its docstring) - this entry point only
    surfaces report_month's own data, matching the one-file-one-month shape
    every other extract_* here returns; the API/UI have no way to act on a
    second report_month from a single upload anyway."""
    suffix = Path(path).suffix.lower()
    if suffix == ".docx":
        return {"enviro": extract_docx(path, report_month, mlabel), "coal": {}}
    if suffix in (".xlsx", ".xlsm"):
        return {"enviro": extract_xlsx(path, report_month, mlabel)[report_month], "coal": {}}
    return extract_pdf(path, report_month, mlabel)


def plant_techno_json(enviro: dict, coal: dict, plant: str) -> dict:
    """techno_data["month"] dict for one plant from one month's extraction
    (SAIL deliberately excluded - see module docstring)."""
    out = {}
    for key, _label, json_key in ENVIRO_PARAM_ORDER:
        v = enviro.get(key, {}).get("month", {}).get(plant)
        if v is not None:
            out[json_key] = v
    cvals = coal.get(plant, {})
    for src_key, json_key in [("pcc", "indigenous_pcc"), ("mcc", "indigenous_mcc"),
                               ("hard", "imported_hard_coal"), ("soft", "imported_soft_coal")]:
        v = cvals.get(src_key)
        if v is not None:
            out[json_key] = v
    return out


def plant_till_techno_json(enviro: dict, plant: str) -> dict:
    """techno_data["till_month"] dict for one plant from one month's
    extraction. Only the .docx path populates "till_month" per param (its
    FY-cumulative column) - the PDF path never does, so this is {} for a
    PDF-sourced extraction (SAIL deliberately excluded, as in
    plant_techno_json)."""
    out = {}
    for key, _label, json_key in ENVIRO_PARAM_ORDER:
        v = enviro.get(key, {}).get("till_month", {}).get(plant)
        if v is not None:
            out[json_key] = v
    return out


def load_folder(folder: str, write: bool = True) -> dict:
    """Extract every "<Mon>'YY.pdf" in folder and (if write=True) merge into
    techno_data (per plant, unit='General') and techno_plan_fy (per plant +
    SAIL, unit='Shop', FY of the LAST month processed - the annual target
    column is FY-constant so any month's PDF carries the same figures).
    Returns {report_month: {"enviro":..., "coal":...}} for inspection either
    way."""
    sys.path.insert(0, str(Path(__file__).parent.parent))
    import db  # noqa: E402

    results = {}
    for path in sorted(list(Path(folder).glob("*.pdf")) + list(Path(folder).glob("*.docx"))):
        parsed = report_month_from_filename(path.name)
        if not parsed:
            continue
        report_month, mlabel = parsed
        results[report_month] = extract_report(str(path), report_month, mlabel)

    if not write:
        return results

    for report_month, blob in sorted(results.items()):
        for plant in PLANTS:
            month_json = plant_techno_json(blob["enviro"], blob["coal"], plant)
            till_json = plant_till_techno_json(blob["enviro"], plant)
            if month_json or till_json:
                db.merge_upsert_techno_data(plant, report_month, "General",
                                             {"month": month_json, "till_month": till_json},
                                             source_file=f"Coal_co2/{Path(folder).name}")

    if results:
        last_month = max(results)
        fy_num = int(last_month[:4]) if int(last_month[5:7]) >= 4 else int(last_month[:4]) - 1
        target_fy = f"{fy_num}-{(fy_num + 1) % 100:02d}"
        enviro = results[last_month]["enviro"]

        for plant in PLANTS:
            plan = db.get_techno_plant_plan(plant, target_fy)
            plan_data = dict(plan.get("data") or {})
            for key, label, _jk in ENVIRO_PARAM_ORDER:
                v = enviro.get(key, {}).get("target", {}).get(plant)
                if v is not None:
                    plan_data[label] = {"value": v, "unit": ENVIRO_KEY_UNITS[_jk]}
            db.save_techno_plant_plan(plant, target_fy, plan_data,
                                       is_user_supplied=plan.get("is_user_supplied", False),
                                       created_by="coal_co2_epi_extractor")

        sail_plan = db.get_sail_techno_plan(target_fy)
        sail_plan_data = dict(sail_plan.get("data") or {})
        for key, label, _jk in ENVIRO_PARAM_ORDER:
            v = enviro.get(key, {}).get("target", {}).get("SAIL")
            if v is not None:
                sail_plan_data[label] = {"value": v, "unit": ENVIRO_KEY_UNITS[_jk]}
        db.save_sail_techno_plan(target_fy, sail_plan_data,
                                  is_user_supplied=sail_plan.get("is_user_supplied", False),
                                  created_by="coal_co2_epi_extractor")

    return results


if __name__ == "__main__":
    import json as _json
    folder_arg = sys.argv[1] if len(sys.argv) > 1 else r"D:\opr-mis1\Report_format\Coal_co2"
    dry_run = "--dry-run" in sys.argv
    res = load_folder(folder_arg, write=not dry_run)
    print(_json.dumps(res, indent=2))
