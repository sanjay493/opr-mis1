"""
Monthly D.O. Letter — "Monthly performance of SAIL" letter due on the 1st of
every month, reporting the PREVIOUS month's performance, plus its two
Annexures (Crude Steel / Finished Steel, plant-wise, each with a Remarks
column). Format-matched against Report_format/DO_Jul-26 Monthly
Performance.docx by mutating a bundled copy of that exact file
(do_letter_templates/do_letter_template.docx) in place — every run/table/
paragraph's own formatting (fonts, borders, bullet numbering, landscape
layout) comes from the template unchanged; only text content changes. See
build_do_letter_docx_bytes().

Also builds the companion "Monthly DO Annexure" workbook (Ministry of Steel
Part A/B/C indicators) the same way — do_annexure_template.xlsx. Only Part A
rows 1-2 (Crude Steel / Finished Steel production, MT) have a source in this
app's DB; everything else (Iron ore/Ferro scrap/PLI/MSME rows, quarterly
capacity utilisation, Part B fund-utilisation, Part C sectoral indicators)
has no data source here and is left exactly as the template shows — blank,
not guessed. See build_do_annexure_xlsx_bytes().

Formulas cross-checked against Jul'26 in the reference docx, reproduced
exactly to the tonne:
  Crude Steel SAIL     = BSP+DSP+RSP+BSL+ISP+ASP+SSP+VISL
                          (page4.py's own "Total Crude Steel" sail_set)
  Finished Steel SAIL* = BSP+DSP+RSP+BSL+ISP+ASP+SSP+VISL + Conversion
                          (page4.py's "SAIL INCL. CONV." row — Conversion is
                          production_table's plant_name='SAIL',
                          item_name='Conversion' row)
Both totals are summed from the UNROUNDED plant figures, then rounded once
at display time — summing already-rounded parts would drift from the
reference by a tonne or two on some months (verified).

Remarks (do_letter_remark_table, keyed report_month/item_name/plant_name,
item_name = 'Crude Steel' | 'Finished Steel') are entered ahead of time via
/api/do-letter/remarks so the letter already has them when generated on the
1st — see main.py.
"""
import os
import io

import db
from constants import FIVE_PLANTS as _5P

_TEMPLATE_DIR = os.path.join(os.path.dirname(__file__), "do_letter_templates")
_DOCX_TEMPLATE = os.path.join(_TEMPLATE_DIR, "do_letter_template.docx")
_XLSX_TEMPLATE = os.path.join(_TEMPLATE_DIR, "do_annexure_template.xlsx")

_MONTHS = ["Jan", "Feb", "Mar", "Apr", "May", "Jun",
           "Jul", "Aug", "Sep", "Oct", "Nov", "Dec"]

# Matches page4.py's own "sail_set" for these two items exactly.
_CRUDE_STEEL_PLANTS = _5P + ["ASP", "SSP", "VISL"]
_FINISHED_STEEL_PLANTS = _5P + ["ASP", "SSP", "VISL"]
_SPECIAL_STEEL_PLANTS = ["ASP", "SSP", "VISL"]

# "Best ever Month production" bullets — item, plant set (matches page4.py's
# sail_set for each). Finished Steel's check is on the plain plant sum, not
# the +Conversion "SAIL*" figure — Conversion isn't tracked far enough back
# to make an all-time-history comparison meaningful.
_BEST_EVER_CHECKS = [
    ("Hot Metal", "Hot Metal", _5P + ["VISL"]),
    ("Saleable Steel", "Saleable Steel", _5P + ["ASP", "SSP", "VISL"]),
    ("Finished Steel", "Finished Steel", _5P + ["ASP", "SSP", "VISL"]),
]


def _mlabel(ym: str) -> str:
    y, m = int(ym[:4]), int(ym[5:7])
    return f"{_MONTHS[m - 1]}'{str(y)[2:]}"


def _prev_month(ym: str) -> str:
    y, m = int(ym[:4]), int(ym[5:7])
    return f"{y - 1}-12" if m == 1 else f"{y}-{m - 1:02d}"


def _gr(cur_v, prev_v):
    if cur_v is None or prev_v is None or prev_v == 0:
        return None
    return round((cur_v - prev_v) / prev_v * 100, 1)


def _fetch_item_month(cur, item: str, month: str, plants: list) -> dict:
    """{plant: value} for one item/month — only plants with a non-NULL value."""
    phs = ",".join("?" for _ in plants)
    cur.execute(
        f"SELECT plant_name, month_actual FROM production_table "
        f"WHERE item_name=? AND report_month=? AND plant_name IN ({phs})",
        [item, month] + plants,
    )
    return {p: v for p, v in cur.fetchall() if v is not None}


def _fetch_conversion(cur, month: str):
    cur.execute(
        "SELECT month_actual FROM production_table "
        "WHERE report_month=? AND plant_name='SAIL' AND item_name='Conversion'",
        (month,),
    )
    r = cur.fetchone()
    return r[0] if r and r[0] is not None else None


def _sum_or_none(vals: dict, plants: list):
    present = [vals[p] for p in plants if p in vals]
    return sum(present) if present else None


def _fetch_remarks(cur, report_month: str, item_name: str) -> dict:
    cur.execute(
        "SELECT plant_name, remark FROM do_letter_remark_table "
        "WHERE report_month=? AND item_name=?",
        (report_month, item_name),
    )
    return {p: r for p, r in cur.fetchall() if r}


def _build_plant_table(cur, report_month, cply_month, prev_month, item, plant_row_order,
                        aggregate_rows, remarks):
    """Rows for one Annexure table. aggregate_rows: list of
    (label, plants, add_conversion) — a subtotal/total row summed from
    `plants`, e.g. ('SAIL', _CRUDE_STEEL_PLANTS, False)."""
    cur_vals  = _fetch_item_month(cur, item, report_month, _CRUDE_STEEL_PLANTS if item == "Total Crude Steel" else _FINISHED_STEEL_PLANTS)
    cply_vals = _fetch_item_month(cur, item, cply_month,   _CRUDE_STEEL_PLANTS if item == "Total Crude Steel" else _FINISHED_STEEL_PLANTS)
    prev_vals = _fetch_item_month(cur, item, prev_month,   _CRUDE_STEEL_PLANTS if item == "Total Crude Steel" else _FINISHED_STEEL_PLANTS)
    conv_cur  = _fetch_conversion(cur, report_month) if item == "Finished Steel" else None
    conv_cply = _fetch_conversion(cur, cply_month) if item == "Finished Steel" else None
    conv_prev = _fetch_conversion(cur, prev_month) if item == "Finished Steel" else None

    def block(plants, add_conv):
        c = _sum_or_none(cur_vals, plants)
        y = _sum_or_none(cply_vals, plants)
        p = _sum_or_none(prev_vals, plants)
        if add_conv:
            c = (c or 0.0) + (conv_cur or 0.0) if (c is not None or conv_cur is not None) else None
            y = (y or 0.0) + (conv_cply or 0.0) if (y is not None or conv_cply is not None) else None
            p = (p or 0.0) + (conv_prev or 0.0) if (p is not None or conv_prev is not None) else None
        return c, y, p

    rows = []
    for plant in plant_row_order:
        c, y, p = block([plant], False)
        rows.append({
            "plant": plant, "cur": c, "cply": y, "prev": p,
            "gr_cply": _gr(c, y), "gr_prev": _gr(c, p),
            "remark": remarks.get(plant, ""),
        })
    for label, plants, add_conv in aggregate_rows:
        c, y, p = block(plants, add_conv)
        # A single-plant "aggregate" (ASP/SSP get their own row in the Crude
        # Steel table, via this same loop) still carries a real remark —
        # only true multi-plant subtotals (5ISPs, SAIL, Special Steel
        # Plants, SAIL*) never do, matching the reference letter.
        remark = remarks.get(plants[0], "") if len(plants) == 1 else ""
        rows.append({
            "plant": label, "cur": c, "cply": y, "prev": p,
            "gr_cply": _gr(c, y), "gr_prev": _gr(c, p),
            "remark": remark,
        })
    return rows


def crude_steel_annexure_rows(cur, report_month, cply_month, prev_month):
    """BSP/DSP/RSP/BSL/ISP/5ISPs/ASP/SSP/SAIL — matches Annexure-A exactly."""
    remarks = _fetch_remarks(cur, report_month, "Crude Steel")
    return _build_plant_table(
        cur, report_month, cply_month, prev_month, "Total Crude Steel",
        plant_row_order=list(_5P),
        aggregate_rows=[
            ("5ISPs", _5P, False),
            *[(p, [p], False) for p in ("ASP", "SSP")],
            ("SAIL", _CRUDE_STEEL_PLANTS, False),
        ],
        remarks=remarks,
    )


def finished_steel_annexure_rows(cur, report_month, cply_month, prev_month):
    """BSP/DSP/RSP/BSL/ISP/5 ISP's/Special Steel Plants/SAIL* — matches
    Annexure-B exactly (ASP+SSP+VISL combined into one row here, unlike
    Crude Steel's table where ASP/SSP get their own rows)."""
    remarks = _fetch_remarks(cur, report_month, "Finished Steel")
    return _build_plant_table(
        cur, report_month, cply_month, prev_month, "Finished Steel",
        plant_row_order=list(_5P),
        aggregate_rows=[
            ("5 ISP's", _5P, False),
            ("Special Steel Plants", _SPECIAL_STEEL_PLANTS, False),
            ("SAIL*", _FINISHED_STEEL_PLANTS, True),
        ],
        remarks=remarks,
    )


def _all_time_totals(cur, item, plants):
    """{report_month: total} across every month in DB history where EVERY
    plant in `plants` has a value — a month with any plant missing is
    excluded rather than compared as a (potentially much lower) partial
    sum, so an incomplete historical month can never falsely look like a
    record or knock a real record off the podium."""
    phs = ",".join("?" for _ in plants)
    cur.execute(
        f"SELECT report_month, plant_name, month_actual FROM production_table "
        f"WHERE item_name=? AND plant_name IN ({phs})",
        [item] + plants,
    )
    by_month = {}
    for rm, p, v in cur.fetchall():
        if v is not None:
            by_month.setdefault(rm, {})[p] = v
    return {rm: sum(vals.values()) for rm, vals in by_month.items() if len(vals) == len(plants)}


def _best_ever(cur, item, plants, report_month):
    """None if report_month isn't a top-2 all-time month for this item;
    otherwise a dict describing the record for the DOCX bullet text."""
    totals = _all_time_totals(cur, item, plants)
    if report_month not in totals:
        return None
    cur_val = totals[report_month]
    others = sorted(((rm, v) for rm, v in totals.items() if rm != report_month),
                     key=lambda kv: -kv[1])
    if not others:
        return {"rank": 1, "value": cur_val, "prev_best": None, "prev_best_month": None}
    best_month, best_val = others[0]
    if cur_val > best_val:
        return {"rank": 1, "value": cur_val, "prev_best": best_val, "prev_best_month": best_month}
    if len(others) > 1 and cur_val > others[1][1]:
        second_month, second_val = others[1]
        return {
            "rank": 2, "value": cur_val,
            "prev_best": best_val, "prev_best_month": best_month,
            "prev_second": second_val, "prev_second_month": second_month,
        }
    return None


def best_ever_bullets(cur, report_month):
    """Narrative bullet strings, e.g. 'Hot Metal production at 1.777 MT
    ( Previous best 1.770 MT in Jul'23)' — same wording pattern as the
    reference letter. Empty list on a month with no new top-2 record."""
    bullets = []
    for label, item, plants in _BEST_EVER_CHECKS:
        r = _best_ever(cur, item, plants, report_month)
        if not r:
            continue
        mt = round(r["value"] / 1000.0, 3)
        prev_mt = round(r["prev_best"] / 1000.0, 3)
        prev_mon = _mlabel(r["prev_best_month"])
        text = f"{label} production at {mt} MT ( Previous best {prev_mt} MT in {prev_mon})"
        if r["rank"] == 2:
            second_mt = round(r["prev_second"] / 1000.0, 3)
            second_mon = _mlabel(r["prev_second_month"])
            text += (f", it is also the 2nd best ever month "
                     f"(previous 2nd best {second_mt} MT in {second_mon})")
        bullets.append(text)
    return bullets


def generate_do_letter_data(report_month: str) -> dict:
    """Everything the DOCX/XLSX builders need for one report_month —
    figures in '000T unless noted. No DB writes."""
    cply_month = db.get_cply_month(report_month)
    prev_month = _prev_month(report_month)
    ytd_months = db.get_ytd_months(report_month)
    ytd_cply_months = [db.get_cply_month(m) for m in ytd_months]

    conn = db.connect()
    cur = conn.cursor()
    try:
        cs_rows = crude_steel_annexure_rows(cur, report_month, cply_month, prev_month)
        fs_rows = finished_steel_annexure_rows(cur, report_month, cply_month, prev_month)
        bullets = best_ever_bullets(cur, report_month)

        cs_sail = next(r for r in cs_rows if r["plant"] == "SAIL")
        fs_sail = next(r for r in fs_rows if r["plant"] == "SAIL*")

        # FY-cumulative (Apr..report_month) SAIL totals, for the XLSX Part A
        # "FY <this> (upto <month>)" / "FY <prior> (upto <month>)" columns.
        def ytd_sum(months, item, plants, add_conv):
            total, found = 0.0, False
            for m in months:
                vals = _fetch_item_month(cur, item, m, plants)
                v = _sum_or_none(vals, plants)
                if add_conv:
                    c = _fetch_conversion(cur, m)
                    v = (v or 0.0) + (c or 0.0) if (v is not None or c is not None) else None
                if v is not None:
                    total += v
                    found = True
            return round(total, 3) if found else None

        cs_ytd      = ytd_sum(ytd_months,      "Total Crude Steel", _CRUDE_STEEL_PLANTS, False)
        cs_ytd_cply = ytd_sum(ytd_cply_months,  "Total Crude Steel", _CRUDE_STEEL_PLANTS, False)
        fs_ytd      = ytd_sum(ytd_months,      "Finished Steel", _FINISHED_STEEL_PLANTS, True)
        fs_ytd_cply = ytd_sum(ytd_cply_months,  "Finished Steel", _FINISHED_STEEL_PLANTS, True)
    finally:
        conn.close()

    return {
        "report_month": report_month, "cply_month": cply_month, "prev_month": prev_month,
        "ytd_months": ytd_months,
        "crude_steel_rows": cs_rows, "finished_steel_rows": fs_rows,
        "crude_steel_sail": cs_sail, "finished_steel_sail": fs_sail,
        "best_ever_bullets": bullets,
        "cs_ytd": cs_ytd, "cs_ytd_cply": cs_ytd_cply,
        "fs_ytd": fs_ytd, "fs_ytd_cply": fs_ytd_cply,
    }


# ---------------------------------------------------------------------------
# DOCX builder — mutates a copy of do_letter_template.docx in place, text
# only. See module docstring for why (fixed row/paragraph counts per month,
# except the "best ever" bullets, which vary 0-3+).
# ---------------------------------------------------------------------------

_TITLE_IDX = 1
_CRUDE_NARRATIVE_IDX = 3
_FINISHED_NARRATIVE_IDX = 5
_DEV_HEADER_IDX = 7
_BEST_HEADER_IDX = 9
_BEST_BULLET_START_IDX = 10
_N_TEMPLATE_BULLETS = 3


def _dir_word(v: float) -> str:
    return "increased" if v >= 0 else "decreased"


def _narrative(item_label, report_label, mt, gr_prev, prev_label, gr_cply, cply_label, annexure_letter):
    return (
        f"The total {item_label} production of SAIL during {report_label} was {mt} MT, "
        f"which has {_dir_word(gr_prev)} by {abs(gr_prev)}% compared to {prev_label} and "
        f"{_dir_word(gr_cply)} by {abs(gr_cply)}% over {cply_label}. "
        f"Details are placed at Annexure {annexure_letter}."
    )


def _set_paragraph_text(paragraph, text: str) -> None:
    """Overwrite a paragraph's visible text while keeping its first run's
    formatting (font/bold/size) — python-docx's own `paragraph.text = ...`
    doesn't exist, and Cell.text-style helpers drop formatting, hence this."""
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(text)
        return
    runs[0].text = text
    for r in runs[1:]:
        r.text = ""


def _remove_paragraph(paragraph) -> None:
    el = paragraph._p
    el.getparent().remove(el)


def _clone_paragraph_after(ref_paragraph, text: str):
    """Deep-copies ref_paragraph's XML (preserving its run/paragraph
    formatting and bullet numbering), inserts the copy immediately after it,
    and sets the copy's text. Used when more "best ever" bullets are needed
    than the template's original 3."""
    import copy
    new_p_el = copy.deepcopy(ref_paragraph._p)
    ref_paragraph._p.addnext(new_p_el)
    from docx.text.paragraph import Paragraph
    new_paragraph = Paragraph(new_p_el, ref_paragraph._parent)
    _set_paragraph_text(new_paragraph, text)
    return new_paragraph


def _fmt_val(v):
    return "" if v is None else str(round(v))


def _fmt_pct(v):
    return "" if v is None else f"{v:.1f}"


def _set_cell_text(cell, text: str) -> None:
    """Multi-line remark cells have one <w:p> per original line — clearing
    only paragraphs[0] (often itself blank, the source of a stray leading
    line break) left every later paragraph's old template text in place.
    Extra paragraphs are removed outright (not just blanked) so a short or
    empty remark doesn't leave trailing blank lines from the old text's
    longer line count."""
    paras = cell.paragraphs
    _set_paragraph_text(paras[0], text)
    for p in paras[1:]:
        _remove_paragraph(p)


def _fill_annexure_table(table, rows, data_start_row: int):
    """rows: list from crude_steel_annexure_rows()/finished_steel_annexure_rows().
    data_start_row: 1 for Table A (row 0 is the column-header row), 2 for
    Table B (row 0 is the merged 'Finished Steel' title, row 1 the header)."""
    for i, r in enumerate(rows):
        row_cells = table.rows[data_start_row + i].cells
        _set_cell_text(row_cells[0], r["plant"])
        _set_cell_text(row_cells[1], _fmt_val(r["cur"]))
        _set_cell_text(row_cells[2], _fmt_val(r["cply"]))
        _set_cell_text(row_cells[3], _fmt_val(r["prev"]))
        _set_cell_text(row_cells[4], _fmt_pct(r["gr_cply"]))
        _set_cell_text(row_cells[5], _fmt_pct(r["gr_prev"]))
        _set_cell_text(row_cells[6], r["remark"])


def build_do_letter_docx_bytes(report_month: str) -> bytes:
    """Renders the Monthly D.O. Letter for report_month, format-matched
    against do_letter_templates/do_letter_template.docx exactly (fonts,
    borders, bullet numbering, landscape layout all come from the template
    unchanged — only text content is replaced)."""
    import docx

    data = generate_do_letter_data(report_month)
    report_label = _mlabel(report_month)
    cply_label = _mlabel(data["cply_month"])
    prev_label = _mlabel(data["prev_month"])

    doc = docx.Document(_DOCX_TEMPLATE)
    paras = doc.paragraphs

    _set_paragraph_text(paras[_TITLE_IDX], f"Monthly performance of SAIL: {report_label}")

    cs = data["crude_steel_sail"]
    _set_paragraph_text(paras[_CRUDE_NARRATIVE_IDX], _narrative(
        "Crude Steel", report_label, round(cs["cur"] / 1000.0, 3),
        cs["gr_prev"], prev_label, cs["gr_cply"], cply_label, "A",
    ))

    fs = data["finished_steel_sail"]
    _set_paragraph_text(paras[_FINISHED_NARRATIVE_IDX], _narrative(
        "Finished Steel", report_label, round(fs["cur"] / 1000.0, 3),
        fs["gr_prev"], prev_label, fs["gr_cply"], cply_label, "B",
    ))

    # The template's own wording named the PRIOR month here even though the
    # rest of the letter reports report_month — kept consistent with
    # report_month instead of reproducing what reads as a drafting slip.
    _set_paragraph_text(paras[_DEV_HEADER_IDX],
                        f"Information on Significant developments during {report_label}:")

    bullets = data["best_ever_bullets"]
    if not bullets:
        _remove_paragraph(paras[_BEST_HEADER_IDX])
        for i in range(_N_TEMPLATE_BULLETS):
            _remove_paragraph(paras[_BEST_BULLET_START_IDX + i])
    else:
        for i in range(_N_TEMPLATE_BULLETS):
            p = paras[_BEST_BULLET_START_IDX + i]
            if i < len(bullets):
                _set_paragraph_text(p, bullets[i])
            else:
                _remove_paragraph(p)
        if len(bullets) > _N_TEMPLATE_BULLETS:
            ref = paras[_BEST_BULLET_START_IDX + _N_TEMPLATE_BULLETS - 1]
            for extra in bullets[_N_TEMPLATE_BULLETS:]:
                ref = _clone_paragraph_after(ref, extra)

    _fill_annexure_table(doc.tables[0], data["crude_steel_rows"], data_start_row=1)
    _fill_annexure_table(doc.tables[1], data["finished_steel_rows"], data_start_row=2)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# XLSX builder — mutates a copy of do_annexure_template.xlsx in place. Only
# "Part A" rows 4-5 (Crude Steel / Finished Steel, the two indicators this
# app actually has a data source for) are touched; everything else (rows
# 6-10, the quarterly capacity-utilisation row, Part B, Part C) is left
# exactly as the template shows — no rated-capacity/import-export/scheme-
# fund/quality-reject data exists in this app's DB, so nothing to compute.
# ---------------------------------------------------------------------------

import calendar as _calendar


def build_do_annexure_xlsx_bytes(report_month: str) -> bytes:
    """Renders the Monthly DO Annexure workbook for report_month, format-
    matched against do_letter_templates/do_annexure_template.xlsx exactly."""
    import openpyxl
    from datetime import date as _date

    data = generate_do_letter_data(report_month)
    y, m = int(report_month[:4]), int(report_month[5:7])
    cply_y = y - 1
    fy_start = y if m >= 4 else y - 1
    last_day = _calendar.monthrange(y, m)[1]
    month_abbr = _MONTHS[m - 1]

    wb = openpyxl.load_workbook(_XLSX_TEMPLATE)
    ws = wb["Part A"]

    py, pm = int(data["prev_month"][:4]), int(data["prev_month"][5:7])
    ws.cell(row=3, column=3, value=_date(py, pm, 1))
    ws.cell(row=3, column=4, value=_date(y, m, 1))
    ws.cell(row=3, column=5, value=_date(cply_y, m, 1))
    ws.cell(row=3, column=6, value=f"FY {fy_start}-{str(fy_start + 1)[2:]} (upto {last_day} {month_abbr})")
    ws.cell(row=3, column=7, value=f"FY {fy_start - 1}-{str(fy_start)[2:]} (upto {last_day} {month_abbr})")
    ws.cell(row=3, column=8, value=f"% Change in Apr-{month_abbr}'{str(y)[2:]} w.r.t. Apr-{month_abbr}'{str(cply_y)[2:]}")

    def _mt(v):
        return round(v / 1000.0, 3) if v is not None else None

    cs, fs = data["crude_steel_sail"], data["finished_steel_sail"]
    ws.cell(row=4, column=3, value=_mt(cs["prev"]))
    ws.cell(row=4, column=4, value=_mt(cs["cur"]))
    ws.cell(row=4, column=5, value=_mt(cs["cply"]))
    ws.cell(row=4, column=6, value=_mt(data["cs_ytd"]))
    ws.cell(row=4, column=7, value=_mt(data["cs_ytd_cply"]))
    # column 8 keeps the template's own "=(F-G)*100/G" formula — recalculated
    # by Excel/LibreOffice on open, same as the '1 page report' generator.

    ws.cell(row=5, column=3, value=_mt(fs["prev"]))
    ws.cell(row=5, column=4, value=_mt(fs["cur"]))
    ws.cell(row=5, column=5, value=_mt(fs["cply"]))
    ws.cell(row=5, column=6, value=_mt(data["fs_ytd"]))
    ws.cell(row=5, column=7, value=_mt(data["fs_ytd_cply"]))

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()
